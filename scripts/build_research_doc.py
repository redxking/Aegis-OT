"""Build the single controlled Aegis-OT research study document."""

from __future__ import annotations

import json
import re
import tempfile
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "research" / "Aegis-OT_Research_Study.docx"
REVISION = json.loads((ROOT / "research" / "revision_log.json").read_text())
MANIFEST = json.loads(
    (ROOT / "results" / "m2-independent-oracle" / "manifest.json").read_text()
)
M3_MANIFEST = json.loads(
    (ROOT / "results" / "m3-physical-modbus" / "manifest.json").read_text()
)
M3_REPRODUCTION_MANIFEST = json.loads(
    (ROOT / "results" / "m3-physical-modbus-reproduction" / "manifest.json").read_text()
)
M3_SUMMARY = json.loads(
    (ROOT / "results" / "m3-physical-modbus" / "summary.json").read_text()
)
PUBLIC_DEMO_EVIDENCE = json.loads(
    (ROOT / "src" / "aegis_ot" / "web_demo" / "evidence.json").read_text()
)
PUBLIC_M3_VERIFICATION = PUBLIC_DEMO_EVIDENCE["m3"]["verification"]
FORMAL_MANIFEST = json.loads(
    (ROOT / "results" / "formal" / "m1-authorization-conformance" / "manifest.json").read_text()
)
FORMAL_INTENDED = next(case for case in FORMAL_MANIFEST["cases"] if case["name"] == "intended")
CURRENT_REVISION = REVISION["revisions"][-1]
PUBLIC_DEMO_REVISION = next(
    revision
    for revision in REVISION["revisions"]
    if revision.get("evidence_artifact") == "src/aegis_ot/web_demo/evidence.json"
)
REVISION_NUMBER = CURRENT_REVISION["revision"]
BLUE = RGBColor(23, 105, 170)
DARK = RGBColor(16, 42, 67)
MUTED = RGBColor(82, 96, 109)
LIGHT_FILL = "EAF2F8"
EXTENDED_PROPERTIES_NAMESPACE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
)


def set_font(run, size: float, *, bold: bool = False, color: RGBColor = DARK) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_alt_text(inline_shape, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description.split(".")[0])


def add_paragraph(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.22
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_font(lead, 10.5, bold=True)
        body = paragraph.add_run(text[len(bold_lead) :])
        set_font(body, 10.5)
    else:
        run = paragraph.add_run(text)
        set_font(run, 10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.2
        set_font(paragraph.add_run(item), 10.5)


def heading(
    doc: Document,
    text: str,
    level: int = 1,
    *,
    page_break_before: bool = False,
) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.paragraph_format.line_spacing = 1.05


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = tbl.rows[0].cells[index]
        cell.width = Inches(widths[index])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(cell, LIGHT_FILL)
        set_cell_margins(cell)
        set_font(cell.paragraphs[0].add_run(header), 9, bold=True)
    set_repeat_table_header(tbl.rows[0])
    prevent_row_split(tbl.rows[0])
    for values in rows:
        row = tbl.add_row()
        prevent_row_split(row)
        cells = row.cells
        for index, value in enumerate(values):
            cells[index].width = Inches(widths[index])
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[index])
            cells[index].paragraphs[0].paragraph_format.space_after = Pt(0)
            set_font(cells[index].paragraphs[0].add_run(value), 8.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def figure(doc: Document, filename: str, caption: str, alt: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(ROOT / "assets" / filename), width=Inches(6.25))
    set_alt_text(shape, alt)
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.keep_with_next = True
    run = caption_p.add_run(caption)
    set_font(run, 9, bold=True, color=MUTED)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.22
    for style_name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 11.5, 8, 4, DARK),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_font(
            header.add_run(
                f"AEGIS-OT-STUDY-001 | CONTROLLED RESEARCH STUDY | REV {REVISION_NUMBER}"
            ),
            8,
            color=MUTED,
        )
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("CONTROLLED RESEARCH STUDY"), 11, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_font(title.add_run("Aegis-OT"), 30, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    set_font(
        subtitle.add_run(
            "Assured Agentic AI for Critical Infrastructure:\nIdentity-Bound Runtime Authorization and Operate-Through-Compromise Resilience"
        ),
        15,
        color=BLUE,
    )
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(author.add_run("Angelis Pseftis"), 12, bold=True)
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        metadata.add_run(
            f"Document AEGIS-OT-STUDY-001 | Revision {REVISION_NUMBER} | 24 August 2026"
        ),
        10,
        color=MUTED,
    )
    boundary = doc.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    boundary.paragraph_format.space_before = Pt(28)
    set_font(boundary.add_run("RESEARCH USE ONLY - SYNTHETIC AND SIMULATED ENVIRONMENTS"), 10, bold=True, color=RGBColor(166, 27, 27))
    doc.add_page_break()


def patch_extended_properties(path: Path, pages: int, words: int) -> None:
    """Preserve useful publication metadata omitted by python-docx."""
    with zipfile.ZipFile(path) as source:
        app_xml = source.read("docProps/app.xml")
        root = ElementTree.fromstring(app_xml)  # noqa: S314 - self-generated DOCX part
        for name, value in (("Pages", pages), ("Words", words)):
            node = root.find(f"{{{EXTENDED_PROPERTIES_NAMESPACE}}}{name}")
            if node is not None:
                node.text = str(value)
        updated = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
        with tempfile.NamedTemporaryFile(dir=path.parent, suffix=".docx", delete=False) as handle:
            temporary_path = Path(handle.name)
        try:
            with zipfile.ZipFile(temporary_path, "w", zipfile.ZIP_DEFLATED) as target:
                for item in source.infolist():
                    target.writestr(
                        item,
                        updated if item.filename == "docProps/app.xml" else source.read(item.filename),
                    )
            # Keep the canonical manuscript's filesystem identity and creation
            # timestamp while replacing its ZIP payload in place.
            path.write_bytes(temporary_path.read_bytes())
        finally:
            temporary_path.unlink(missing_ok=True)


def build() -> None:
    doc = Document()
    configure_styles(doc)
    props = doc.core_properties
    props.author = "Angelis Pseftis"
    props.last_modified_by = "Angelis Pseftis"
    props.title = "Aegis-OT Research Study"
    props.subject = "Identity-bound runtime assurance for simulated critical infrastructure"
    props.keywords = "Aegis-OT, runtime assurance, workload identity, delegation, OT security"
    props.comments = "Controlled research document. Author and editor: Angelis Pseftis."
    props.created = datetime.fromisoformat(REVISION["created_at_utc"].replace("Z", "+00:00"))
    props.modified = datetime.now(UTC)
    props.revision = len(REVISION["revisions"])
    add_cover(doc)

    heading(doc, "Document Control")
    table(
        doc,
        ["Field", "Controlled value"],
        [
            ["Document identifier", "AEGIS-OT-STUDY-001"],
            ["Revision", REVISION_NUMBER],
            ["Author and editor", "Angelis Pseftis"],
            ["Canonical file", "research/Aegis-OT_Research_Study.docx"],
            ["Evidence boundary", "Reconstructed repository; lost historical artifacts are not current evidence"],
        ],
        [1.65, 4.85],
    )
    add_paragraph(doc, "Control statement: This DOCX is the single authoritative editable manuscript. Git history and research/revision_log.json preserve its revision trail. Supporting Markdown, figures, renders, and data are not alternate manuscripts.")
    doc.add_page_break()
    heading(doc, "Revision History", 2)
    revision_rows = []
    for revision in REVISION["revisions"]:
        evidence_link = (
            revision.get("evidence_label")
            or revision.get("experiment_manifest")
            or revision.get("evidence_artifact")
            or "No linked artifact"
        )
        revision_rows.append(
            [
                revision["revision"],
                revision["timestamp_utc"],
                revision["editor"],
                revision["description"],
                evidence_link,
            ]
        )
    table(
        doc,
        ["Revision", "UTC timestamp", "Editor", "Substantive change", "Evidence link"],
        revision_rows,
        [0.55, 1.15, 0.9, 2.7, 1.2],
    )
    heading(doc, "Contents", page_break_before=True)
    sections = [
        "1 Abstract", "2 Executive Summary", "3 Introduction and Research Gap", "4 Research Objectives, Questions, and Hypotheses",
        "5 Scope, Assumptions, and Claim Boundaries", "6 Threat Model", "7 System Requirements", "8 Reference Architecture",
        "9 Authorization and Safety Model", "10 Formal Specification Strategy", "11 Experimental Methodology",
        "12 Implementation and Verification Status", "13 M2 Controlled Results", "14 M3 Physical Process-Boundary Results",
        "15 Multi-VM Integration Plan", "16 Operate-Through-Compromise Test Program", "17 Scale and Economic Analysis Plan",
        "18 Reproducibility and Data Management", "19 Risks, Limitations, and Validity Threats",
        "20 Work Plan and Completion Criteria", "21 Conclusions", "22 References", "23 Requirements Traceability Appendix",
        "24 Data and Message Schemas Appendix", "25 Reproduction Runbook Appendix", "26 Version-Control Audit Trail Appendix",
    ]
    add_bullets(doc, sections)
    doc.add_page_break()

    heading(doc, "1. Abstract")
    add_paragraph(doc, "Aegis-OT investigates whether an independently enforced runtime-assurance control plane can prevent a validly authenticated but compromised, misled, stale, or faulty AI agent from executing actions outside its delegated mission and modeled cyber-physical safety envelope while preserving legitimate containment and recovery. The reconstructed implementation binds typed proposals to Ed25519-signed delegation chains, contextual policy, state freshness, replay protection, modeled safety, one-time execution permits, signed command acknowledgments, readback, and hash-chained evidence. M2 separates the runtime kernel from a decimal-arithmetic reference model and exposes a material guardband-sensitivity gap. M3 adds a pandapower CIGRE MV steady-state plant and a permit-aware PyModbus virtual device in a spawned process. In 30 fixed-condition sessions, the retained primary M3 run observed no non-nominal modeled effects, 30 of 30 nominal completions, no replay effect, no fail-fast unknown-effect outcome, and a complete narrow proposal/decision/terminal-hash indicator; a local rerun reproduced the timing-independent outcome hash. M4a adds a capability-separated deterministic-local path in which the plant, signed observer, and Python research virtual PLC run in distinct same-host processes, the controller has no plant-apply handle, and completion requires one valid applied acknowledgment plus a directly transaction-linked signed post observation. M4a is locally conformance-tested but has no retained or offline-verifiable experiment package. These findings establish bounded local conformance, not physical accuracy, independent replication, field effectiveness, or production readiness.")

    heading(doc, "2. Executive Summary")
    add_paragraph(doc, "The central engineering decision is that identity is necessary but insufficient. An agent never receives direct simulated-control authority; it receives authority to propose a bounded action. The gateway remains the policy-enforcement point and fails closed when identity, delegation, state, replay, policy, safety, or required approval cannot be established.")
    add_bullets(doc, [
        "Current verified implementation evidence: an isolated candidate tree passed 478 tests with 92.05 percent branch-aware coverage on the local Python 3.14 host; repository linting, strict mypy, and schema-drift checks were clean. The candidate used committed retained-evidence files while user-modified result files remained untouched in the primary working tree. Coverage is implementation-conformance evidence, not independent validation.",
        f"Current formal evidence: TLC explored {FORMAL_INTENDED['states_generated']:,} generated and {FORMAL_INTENDED['distinct_states']:,} distinct states to depth {FORMAL_INTENDED['search_depth']} in the intended bounded configuration without an invariant or liveness violation. Sixteen targeted weakened configurations each produced the expected counterexample. This is bounded model evidence, not proof of the Python implementation or physical process.",
        "M2 evidence: 8,640 decision records across eight baselines, 12 reviewed scenario templates, 30 master seeds, and 36 sampled trials per seed per baseline. A separate rerun matched deterministic outcome hash b5ad54a6984f659f961975adaf386eade41f733307c289ea7c3ecaa11c6b5b90.",
        "M2 result: the assured path produced zero unauthorized executions among 450 authorization-negative records but executed 60 percent of the 450 reference-unsafe records because three conservative guardband templates remained inside the runtime kernel limits. This is a threshold-sensitivity finding, not physical evidence.",
        f"M3 evidence: the primary manifest records 30 fresh child-process sessions, 150 fixed-condition trials, 270 chained evidence events, commit {M3_MANIFEST['git']['commit'][:7]}, and a clean-start flag. The controlled package and local rerun pass the offline verifier from the matching clean checkout and share deterministic outcome SHA-256 {M3_MANIFEST['deterministic_outcome_sha256']}. The unsigned historical metadata is not external attestation.",
        "M3 result: 0 of 120 primary-package non-nominal trials produced a modeled effect or unauthorized device application under the registered end-to-end metric, whose denominator includes 60 gateway no-dispatch cases; all 30 nominal commands were applied, acknowledged, and read back; and 0 of 30 replays produced a second effect. The fail-fast runner observed no unknown-effect outcome, and the narrow proposal/decision/terminal-hash indicator was complete for 150 of 150 trials. These are fixed-condition conformance checks from one deterministic model and host, not device-dispatch-conditional, ambiguity-rate, or field-rate estimates.",
        "M4a implementation milestone: the authoritative plant, signed observer, and Python research virtual PLC run in distinct spawned processes with separate observer and PLC keys and boot epochs. The closed-loop controller has observe, simulate, and dispatch ports but no plant-apply handle; the trusted coordinator harness separately retains lifecycle administration and permit-signing authority. One local smoke transaction completed with one dispatch and zero retries, and adversarial tests cover the six terminal effect-certainty states. The smoke artifacts and replay ledger are not retained, so this is not an experiment, reproduction, independent observer result, segmented deployment, or WP4 exit.",
        "Public demonstration: the packaged read-only explorer, evidence-builder check, browser surface, isolated wheel, and container surface were verified locally. These checks establish local implementation, packaging, and presentation behavior only; they are not observed remote CI, deployed isolation, evidence that M3 ran in containers, independent validation, operational validation, or production readiness.",
        "Next decision gate: define and retain an offline-verifiable M4a evidence package and independently operated measurement or model path, then evaluate HELICS coordination, OpenPLC or hardware I/O, durable device identity, segmented networking, post-commit recovery, and externally anchored evidence under separate acceptance gates.",
    ])

    doc.add_page_break()
    heading(doc, "3. Introduction and Research Gap")
    add_paragraph(doc, "Workload identity, delegated authorization, policy engines, runtime assurance, cyber-physical safety analysis, and tamper-evident logging are established fields. Aegis-OT does not claim novelty for those elements individually. The research contribution under evaluation is the integration and reproducible testing of those mechanisms as an independent action-control plane for adversarial autonomous-agent operation in simulated critical infrastructure.")
    add_paragraph(doc, "The present gap statement is provisional. Public evidence has not yet established a mature benchmark that jointly measures cryptographic agent identity, attenuation and revocation, contextual cyber policy, physical-state validation, operate-through-compromise behavior, evidence reconstruction, fleet scale, and governance cost. Absence of identified public evidence is not proof that no comparable system exists; the related-work search must actively seek counterexamples and narrow the claim.")

    heading(doc, "4. Research Objectives, Questions, and Hypotheses")
    add_paragraph(doc, "Primary objective: determine whether independently enforced, identity-bound and state-aware authorization reduces unsafe or unauthorized execution without imposing unacceptable mission, latency, availability, evidence, or operator costs.")
    table(doc, ["ID", "Research question", "Initial hypothesis"], [
        ["RQ1", "Unsafe-action containment", "B3 reduces unsafe execution relative to direct, identity-only, and static-policy baselines in defined scenarios."],
        ["RQ2", "Operational availability", "Bounded recovery authority can preserve safe mission actions with measurable false-block cost."],
        ["RQ3-RQ5", "Blast radius, state value, and revocation", "Attenuation and fresh physical state reduce reachable consequences; revocation remains bounded under stated availability assumptions."],
        ["RQ6-RQ9", "Performance, scale, economics, and evidence", "Assurance costs are measurable and increase with chain depth, state simulation, evidence volume, and approval burden."],
        ["RQ10", "Generalizability", "Some authorization properties transfer across domains; physical findings remain model-specific."],
    ], [0.7, 2.4, 3.4])

    heading(doc, "5. Scope, Assumptions, and Claim Boundaries")
    add_bullets(doc, [
        "In scope: synthetic telemetry, typed proposals, simulated assets, bounded adversarial conditions, local and future multi-node deployment, and reproducible evidence.",
        "Out of scope: production OT, real utility credentials, live third-party infrastructure, sub-cycle protection claims, and operational authorization.",
        "Formal guarantees apply only to the stated model, constants, invariants, trusted computing base, and explored state space.",
        "Empirical findings apply only to recorded versions, configurations, seeds, scenarios, and hosts.",
        "The v0.1 safety thresholds are provisional research parameters, not universal utility settings.",
    ])

    heading(doc, "6. Threat Model")
    add_paragraph(doc, "The adversary may compromise an authenticated agent, possess a bounded credential, forge or amplify a grant, replay a proposal, poison synthetic observations, delay revocation, exploit stale policy or state, create conflicting proposals, or disrupt a supporting service. The gateway, cryptographic implementation, trusted host, and evidence head are initially trusted; later work must reduce and test that trusted computing base.")
    add_paragraph(doc, "M4a refines the local trusted computing base without claiming host isolation. The plant, signed observer, and research virtual PLC use distinct processes, endpoints, signing keys, and boot epochs, and only the plant-spawned PLC child receives the raw apply capability. The controller object has no apply handle, but the enclosing coordinator still holds the permit signer and administrative lifecycle clients. A compromised observer can invalidate observation-origin claims, a compromised PLC possesses apply authority, a compromised plant invalidates both simulated and observed state, and a compromised same-user host can bypass the application-level separation.")
    table(doc, ["Threat family", "Required control", "Primary measure"], [
        ["Identity and credential compromise", "Short-lived identity, scope-bound grant, revocation", "Unauthorized execution and revocation delay"],
        ["Instruction or telemetry manipulation", "Typed proposal, freshness, independent state evaluation", "Unsafe escape and state divergence"],
        ["Replay and concurrency", "Atomic nonce reservation and TOCTOU check", "Duplicate execution and conflict rate"],
        ["Supervisor compromise", "Full-chain attenuation and bounded depth", "Agent and asset blast radius"],
        ["Service loss", "Explicit fail-closed and bounded degraded modes", "Availability, recovery time, residual risk"],
    ], [1.7, 2.9, 1.9])

    heading(doc, "7. System Requirements")
    table(doc, ["Requirement", "Statement", "Verification method"], [
        ["REQ-AUTH-001", "No action executes without verified actor identity and a valid full delegation chain.", "Negative and property tests; formal invariant"],
        ["REQ-DELEG-002", "A child grant cannot exceed parent resources, operations, time, risk, or delegation depth.", "Chain fuzzing and model checking"],
        ["REQ-SAFE-003", "A modeled-unsafe candidate transition is denied by an evaluator outside agent reasoning.", "Scenario tests and independent-oracle comparison"],
        ["REQ-REPLAY-004", "Nonce reuse is atomically rejected within the retention window.", "Concurrency test"],
        ["REQ-EVID-005", "Every gateway decision creates linked evidence sufficient for reconstruction.", "Chain verification and completeness audit"],
        ["REQ-TOCTOU-006", "The adapter rejects authorization when state version changes before execution.", "Integration test"],
        ["REQ-EXEC-007", "A device applies only an exact, signed, unexpired, audience-bound, single-use permit and completion requires signed acknowledgment plus readback.", "M3 device, replay, acknowledgment, and offline-verifier tests"],
        ["REQ-CAP-008", "The controller cannot directly apply plant state; completion requires one PLC dispatch, a valid applied acknowledgment, and a separate transaction-linked signed post observation.", "M4a capability-negative, contract, process, lifecycle, and controller fault-path tests"],
    ], [1.05, 3.65, 1.8])

    heading(doc, "8. Reference Architecture")
    figure(doc, "architecture.png", "Figure 1. Aegis-OT reference action path.", "Architecture diagram showing synthetic telemetry flowing through a bounded agent, typed proposal, Aegis-OT gateway, authorization-bound command adapter, and simulated process. Gateway checks include identity, delegation, policy, freshness, replay, safety, approval, and evidence.")
    add_paragraph(doc, "The reference architecture separates observation, agent, authorization, safety, control, simulation, and evidence responsibilities. The original executable path remains an in-process approximation. M4a adds an application-level same-host topology with distinct plant, signed-observer, and research virtual-PLC processes. Its controller can resolve observations, simulate candidates, and dispatch once, but cannot call plant apply. The plant owns the authoritative pandapower instance and creates the sole apply endpoint for its PLC child. The trusted coordinator still owns the permit signer and lifecycle clients, so this is capability separation, not isolation from a hostile coordinator or host.")

    heading(doc, "9. Authorization and Safety Model")
    add_paragraph(doc, "The execution condition is: Executed(a) implies authenticated actor, valid and attenuated delegation, contextual policy permission, fresh matching state, unique nonce, modeled-safe candidate transition, and required approval. The adapter additionally binds the decision to the proposal and state version at execution time.")
    figure(doc, "decision_sequence.png", "Figure 2. Authorization and execution evidence sequence.", "Sequence diagram with agent, gateway, trust services, safety kernel, adapter, and evidence lanes. It shows proposal submission, identity and delegation verification, candidate transition evaluation, evidence append, command authorization, and acknowledgment.")

    heading(doc, "10. Formal Specification Strategy")
    add_paragraph(doc, f"The expanded TLA+ state machine models submission, authorization or denial, approval, dispatch, acknowledgment, execution, delegation validity, ancestor revocation with an explicit propagation bound, expiry, replay, policy and state consistency, conflicting actions, evidence, compromise, quarantine, and decision liveness. TLC version {FORMAL_MANIFEST['tool_version']} checked the intended configuration at commit {FORMAL_MANIFEST['git_commit'][:7]} with model SHA-256 {FORMAL_MANIFEST['model_sha256']}. It generated {FORMAL_INTENDED['states_generated']:,} states, found {FORMAL_INTENDED['distinct_states']:,} distinct states, reached depth {FORMAL_INTENDED['search_depth']}, and reported no invariant or liveness violation. All sixteen deliberately weakened cases reported their expected invariant violation. The result applies only to this abstraction, configuration, fairness condition, constants, tool build, and explored state space.")
    add_paragraph(doc, "Initial model-check finding: The first TLC run against the earlier scaffold found a specification defect: its revocation invariant retroactively treated a later grant revocation as invalidating an execution that had already completed. Revision 0.3 corrects this by recording whether revocation, expiry, state staleness, acknowledgment absence, or quarantine was effective at execution time. This was a defect in the specification, not evidence of a corresponding runtime exploit.", bold_lead="Initial model-check finding:")
    add_paragraph(doc, "M4a did not extend the committed TLA+ model. Its process topology, canonical-JSON IPC framing, observer and PLC boot epochs, direct transaction pre/post linkage, orderly PLC-child replacement, and transport-ambiguity classifications are covered by runtime implementation tests only. The earlier TLC results must not be represented as formal verification of those mechanisms.")

    heading(doc, "11. Experimental Methodology")
    add_paragraph(doc, "The controlled M2 experiment executes eight baselines and ablations against 12 human-reviewed synthetic scenario templates spanning nominal, consequence-unsafe, guardband, identity, delegation-scope, freshness, confidence, and approval conditions. Thirty master seeds each run three shuffled catalog cycles. Trial seeds sample bounded action parameters within each template; execution stops if either the kernel or reference classification departs from the catalog's reviewed expectation. B0-B3 preserve the original direct, identity-only, static-policy, and complete-gateway paths. B4-B7 add contextual ABAC, risk-aware, safety-without-delegation, and delegation-without-freshness comparisons.")
    add_paragraph(doc, "The reference oracle receives the proposal and pre-action state and independently calculates the candidate state with decimal arithmetic. It does not consume the kernel's predicted state. Its tighter load, voltage, thermal, isolation, and battery guardbands deliberately expose sensitivity. Primary measures are conditional unsafe-action escape, unauthorized execution, false block, mission correctness, decision latency, and kernel-oracle disagreement. Wilson 95 percent intervals characterize records under the balanced synthetic design; they do not capture model-form or field uncertainty.")
    add_paragraph(doc, "The preregistered M3 increment executes five ordered conditions in each of 30 fresh child-process sessions: unknown identity, stale state, a permit whose audience is altered after signing, nominal permitted execution, and reuse of the nominal permit. The parent owns the gateway, command translation, permit issuer, and verified client. The child owns the PyModbus mailbox, permit enforcement, replay state, signed acknowledgments, and authoritative pandapower plant. The tested command isolates one registered line only after the gateway decision, candidate solve, evidence binding, and device-audience checks succeed. The parent accepts completion only after signed acknowledgment verification and post-action readback.")
    add_paragraph(doc, "The audience-alteration fixture changes the audience field after permit issuance. The device checks audience before signature and returns a signed permit_wrong_audience rejection without applying the command. Changing the field also invalidates the permit's original signature, so this condition does not exercise a separately issued, validly signed wrong-audience permit.")
    add_paragraph(doc, "M3 uses pandapower 3.5.4 to instantiate the packaged CIGRE MV network with all distributed-energy resources and runs balanced steady-state Newton-Raphson power flow. Seeds vary process sessions, keys, identifiers, and boot epochs; they do not vary the network, operating point, command, solver, or condition parameters. Wilson intervals describe repeated fixed-condition outcomes only. Candidate assessment, authoritative commit, and readback share the same model and child process, so agreement is not independent physical validation.")
    add_paragraph(doc, "M4a is an implementation-conformance increment rather than a controlled experiment. A fresh observer-signed pre snapshot is resolved and verified before authorization; candidate simulation binds that exact state; a short-lived permit binds the target PLC identity, key, and boot epoch; and the controller makes at most one PLC call. Plant apply atomically compares the authorized state version, state digest, and observation digest. Completion requires a valid PLC-signed applied acknowledgment plus a newly captured observer-signed post snapshot that directly identifies the transaction's pre envelope and matches the permit expectation. That direct predecessor is transaction-local, not a continuous global observation chain. The controller never retries automatically and terminates as not_dispatched, candidate_rejected, plc_rejected, unknown_effect, observation_diverged, or completed.")

    heading(doc, "12. Implementation and Verification Status")
    table(doc, ["Capability", "Current state", "Evidence boundary"], [
        ["Typed proposal and state", "Closed Pydantic models with operation-specific finite parameter validation and generated JSON Schema", "Local conformance only"],
        ["Signed delegation", "Ed25519 full-chain validation and attenuation", "Local keys; no SPIRE integration"],
        ["Gateway", "Fail-closed decision path with replay and freshness", "Single process"],
        ["Safety and oracle", "Separate candidate-state implementations with intentionally different guardbands", "Code-path independence only; no physical validation"],
        ["Physical command path", "Signed one-time permit, PyModbus loopback process boundary, transactional pandapower apply, signed acknowledgment, and readback", "One host and virtual device; no OpenPLC, hardware, or segmented OT network"],
        ["M4a capability-separated path", "Distinct same-host plant, signed-observer, and research virtual-PLC processes; observation-bound compare-and-swap; one dispatch; explicit effect-certainty states", "Coordinator remains privileged; same plant/model; transient evidence only; no hostile-host isolation"],
        ["Evidence", "Hash-chained decisions plus M3 event, permit, acknowledgment, source, schema, configuration, and artifact bindings", "M3 manifest is unsigned and its verification keys are package-internal"],
        ["Verification", "478 tests in the isolated candidate tree; repository ruff clean; strict mypy and schema-drift checks clean; 92.05 percent branch-aware coverage", "Python 3.14 local host; committed evidence copies used; remote CI not observed"],
        ["Public demonstration", "Packaged read-only evidence explorer; public API exposes health and retained-evidence reads only", "Local browser, isolated-wheel, and container checks; no deployment or external validation"],
        ["Formal model", "17 safety/type invariants and one decision-liveness property checked; 16 weakened cases produced expected violations", "Bounded abstraction; not implementation or physical proof"],
    ], [1.55, 2.25, 2.7])

    heading(doc, "13. M2 Controlled Results")
    summary = MANIFEST["summary"]
    result_rows = []
    for baseline, values in summary.items():
        unsafe_ci = values["unsafe_action_escape_ci95"]
        unauthorized_ci = values["unauthorized_execution_ci95"]
        mission_ci = values["mission_success_ci95"]
        result_rows.append([
            baseline.split("_")[0],
            f"{values['unsafe_action_escape_rate']:.0%} [{unsafe_ci['lower']:.1%}, {unsafe_ci['upper']:.1%}]",
            f"{values['unauthorized_execution_rate']:.0%} [{unauthorized_ci['lower']:.1%}, {unauthorized_ci['upper']:.1%}]",
            f"{values['false_block_rate']:.0%}",
            f"{values['mission_success_rate']:.0%} [{mission_ci['lower']:.1%}, {mission_ci['upper']:.1%}]",
        ])
    table(doc, ["Baseline", "Unsafe escape [95% CI]", "Unauthorized [95% CI]", "False block", "Mission correct [95% CI]"], result_rows, [1.05, 1.55, 1.55, 0.85, 1.5])
    figure(doc, "baseline_results.png", "Figure 3. M2 synthetic baseline and ablation outcomes.", "Grouped bar chart comparing conditional unsafe-action escape and unauthorized execution across eight baselines. Each baseline has 1,080 trial records from 30 master seeds and 12 sampled scenario templates.")
    add_paragraph(doc, "Interpretation: B3_ASSURED recorded zero unauthorized executions among 450 authorization-negative records (Wilson 95 percent upper bound 0.85 percent), zero false blocks among 180 authorized reference-safe records (upper bound 2.09 percent), and 75 percent overall mission correctness (72.3-77.5 percent). It nevertheless executed 270 of 450 reference-unsafe records, a 60 percent conditional escape rate (55.4-64.4 percent), because the load, thermal, and voltage guardband templates were inside the runtime kernel's looser thresholds. B6 and B7 matched the 60 percent physical escape rate but each executed 20 percent of authorization-negative records, isolating the value of the omitted delegation or freshness control. Rates are conditional on this balanced catalog and do not estimate operational incident likelihood.")
    add_paragraph(doc, f"Reproducibility: the controlled run started from clean commit {MANIFEST['git_commit'][:7]}, wrote {MANIFEST['total_trial_records']:,} records, and produced deterministic outcome SHA-256 {MANIFEST['deterministic_outcome_sha256']}. A second run produced the same outcome hash. Raw hashes differ as expected because host timing is retained separately. B3 mean in-process decision latency was {summary['B3_ASSURED']['mean_decision_latency_ms']:.3f} ms on the recorded host; this is a local development measurement, not a deployment latency claim.")

    heading(doc, "14. M3 Physical Process-Boundary Results")
    m3_rows = []
    m3_condition_labels = {
        "unknown_identity": "Unknown identity",
        "stale_state": "Stale state",
        "wrong_audience_permit": "Audience altered post-signing",
        "nominal_permitted_execution": "Nominal execution",
        "permit_replay": "Permit replay",
    }
    for condition in (
        "unknown_identity",
        "stale_state",
        "wrong_audience_permit",
        "nominal_permitted_execution",
        "permit_replay",
    ):
        values = M3_SUMMARY["by_condition"][condition]
        m3_rows.append(
            [
                m3_condition_labels[condition],
                str(values["trials"]),
                str(values["state_effects"]),
                str(values["device_applied"]),
                str(values["unknown_effects"]),
            ]
        )
    add_paragraph(doc, f"The operator run record identifies commit {M3_MANIFEST['git']['commit'][:7]} as the prepared clean checkout for the primary M3 run. The unsigned manifest records a clean-start flag and the interval {M3_MANIFEST['started_at_utc']} through {M3_MANIFEST['completed_at_utc']}. The package retains {M3_MANIFEST['session_count']} fresh child-process sessions, {M3_MANIFEST['trial_record_count']} trial records, and {M3_MANIFEST['event_record_count']} chained evidence events. The registered plant model digest is {M3_MANIFEST['model_digest']}. Historical Git, host, and time fields are self-asserted package metadata rather than external attestation.")
    table(
        doc,
        ["Condition", "Trials", "Modeled effects", "Device applied", "Unknown effect"],
        m3_rows,
        [2.2, 0.65, 1.05, 1.0, 1.05],
    )
    figure(
        doc,
        "m3_physical_results.png",
        "Figure 4. M3 fixed-condition conformance outcomes and single-host measured path latency.",
        "M3 results table with 30 trials per condition. Unknown identity, stale state, the permit whose audience was altered after signing, and replay show no modeled effect; nominal execution shows 30 modeled effects and 30 device applications. Box-and-point plots show single-host measured path latency by condition, with explicit non-field-validation caveats.",
    )
    denied_ci = M3_SUMMARY["denied_command_effect_rate_ci95"]
    nominal_ci = M3_SUMMARY["nominal_closed_loop_completion_rate_ci95"]
    unknown_ci = M3_SUMMARY["unknown_effect_rate_ci95"]
    trace_ci = M3_SUMMARY["evidence_trace_completeness_rate_ci95"]
    add_paragraph(doc, f"Primary-package conformance: unknown identity and stale state were denied before dispatch; all 30 post-signing audience-altered permit artifacts received signed permit_wrong_audience device rejections; all 30 nominal commands were applied, signed, and read back; and all 30 replay attempts received signed rejections without a second modeled effect. Because audience was checked before signature, these records characterize the implemented rejection path, not a validly signed wrong-audience permit. Across 120 non-nominal trials, 0 modeled effects and 0 unauthorized device applications were observed under the registered end-to-end metric (two-sided 95 percent Wilson upper bound {denied_ci['upper']:.3%}). That denominator includes 60 gateway no-dispatch trials and is not a device-dispatch-conditional acceptance rate. Nominal completion was 30/30 (lower bound {nominal_ci['lower']:.3%}). Unknown effects were 0/150 (upper bound {unknown_ci['upper']:.3%}), but any unknown effect fails the controlled run; this is a conformance-completeness check, not an ambiguity-rate estimate. The narrow proposal/decision/terminal-hash trace indicator was complete for 150/150 trials (lower bound {trace_ci['lower']:.3%}); the stronger integrity and semantic checks are reported separately by the offline verifier. Zero observations do not establish an impossible event or a field failure rate.")
    nominal_state = M3_SUMMARY["nominal_post_state"]
    nominal_latency = M3_SUMMARY["by_condition"]["nominal_permitted_execution"]["latency"]["end_to_end_ms"]
    add_paragraph(doc, f"For the single deterministic nominal fixture, the post-action minimum voltage was {nominal_state['minimum_voltage_pu']['mean']:.10f} per unit, maximum line loading was {nominal_state['maximum_line_loading_pct']['mean']:.8f} percent, and synthetic priority-load service remained {nominal_state['priority_load_served_pct']['mean']:.0f} percent with no registered voltage, thermal, or supervisory unsafe-state flag. The same physical values in all 30 sessions reflect one fixed model, operating point, and command. Controlled nominal host latency had mean {nominal_latency['mean_ms']:.3f} ms, median {nominal_latency['median_ms']:.3f} ms, and range {nominal_latency['minimum_ms']:.3f}-{nominal_latency['maximum_ms']:.3f} ms; timing is excluded from the deterministic outcome hash and is not an OT performance bound.")
    add_paragraph(doc, f"Evidence interpretation: exactly 90 primary-package trials contain signed device acknowledgments. The 60 gateway no-dispatch trials use the registered verified/not-applicable convention because no device acknowledgment should exist; the reproduction package has the same per-package counts. From the historical matching checkout at commit {M3_MANIFEST['git']['commit'][:7]}, both packages passed all nine offline-verifier checks and reproduced outcome SHA-256 {M3_MANIFEST['deterministic_outcome_sha256']}. The public-demo projection associated with commit {PUBLIC_DEMO_REVISION['git_commit'][:7]} records all {len(PUBLIC_M3_VERIFICATION['internal_checks'])} package-internal checks passing for both retained packages while checkout_bindings reports {PUBLIC_M3_VERIFICATION['current_checkout_binding_status']}, as expected because the current implementation postdates the recorded M3 execution commit. This projection is not a new M3 run, a current matching-checkout execution, or independent replication. The reproduction manifest reports a separate experiment identifier, {M3_REPRODUCTION_MANIFEST['experiment_id']}, and matching source and lock-file hashes, host metadata, Python and selected component versions, model, seed, and conditions. It is local outcome reproduction under matching recorded conditions, not proof of an identical environment or independent replication. The manifests are unsigned and the verification keys are package-internal; verification establishes internal consistency with the matching checkout, not origin, custody, physical-model validity, external authenticity, or field-device identity.")

    heading(doc, "15. Multi-VM Integration Plan")
    add_paragraph(doc, "The six-node scaffold separates management, trust, agent, gateway, OT, and simulation functions. The gateway must become the only routed path between the agent network and the OT environment. Exit evidence includes route and firewall inspection, negative connectivity tests, packet capture, gateway partition behavior, and host-specific deployment documentation. The present Vagrant file is an unvalidated VirtualBox scaffold; it is not evidence that a six-VM range has been deployed.")

    heading(doc, "16. Operate-Through-Compromise Test Program")
    add_bullets(doc, [
        "Compromise a leaf agent while unrelated agents retain bounded mission functions.",
        "Compromise or revoke a supervisor and measure descendant authority and propagation delay.",
        "Remove identity, policy, evidence, or gateway services and measure fail-closed behavior and recovery capacity.",
        "Inject delayed, replayed, biased, or contradictory telemetry and measure state-aware authorization value.",
        "Define quarantine and recovery authority before execution; do not introduce an emergency bypass that recreates direct control.",
    ])

    heading(doc, "17. Scale and Economic Analysis Plan")
    add_paragraph(doc, "Fleet experiments will use logical identities and synthetic authorization workloads at 10, 100, 1,000, and 10,000 agents rather than one VM per agent. Measures include throughput, queue delay, delegation graph complexity, revocation propagation, policy distribution, evidence volume, operator span, incident-response effort, and marginal governance cost. Any cost finding will state labor rates, infrastructure assumptions, utilization, retention, and sensitivity ranges.")

    heading(doc, "18. Reproducibility and Data Management")
    add_paragraph(doc, "The M2 manifest records clean-start Git state, scenario catalog and source hashes, all master seeds, baseline definitions, component versions, host details, raw-data path, timing-inclusive raw SHA-256, timing-independent outcome SHA-256, analyst, and known limitations. Raw JSONL retains sampled parameters and conditional outcomes. The original exploratory run remains separately labeled and is excluded from M2 comparison.")
    add_paragraph(doc, "Each M3 package contains a manifest that records the clean-start flag; source, schema, project, lock, configuration, and artifact hashes; all session seeds; component versions; model digest; summary statistics; deterministic projection; host details; boundary conditions; analyst; and limitations. Separate package artifacts retain per-session process identities and verification keys, benchmark provenance, solver settings, raw trials, and chained events; the manifest binds those artifacts by hash. The offline verifier checks artifact hashes, counts, event chains, trial semantics, deterministic outcomes, summaries, configuration bindings, and matching-checkout bindings. The primary and local reproduction directories are retained separately and must not be overwritten. The retained raw records used for derived figures pass all eight package-internal checks for both packages; checkout binding is reported separately and must not be represented as current-source matching.")
    add_paragraph(doc, "M4a currently has no equivalent retained package. The capability-smoke command reports live process identifiers, health counters, dispatch and retry counts, and an in-memory evidence-chain check, but it does not export the signed observations, permit, PLC acknowledgment, public trust anchors, capability-negative probes, replay provenance, or a manifest. The evidence chain and mode-0600 orderly-restart ledger are removed at normal shutdown. A future M4a evidence gate requires canonical serialization, registered trust anchors, artifact hashes, an offline verifier, unique output directories, and an independently defined replication protocol.")

    heading(doc, "19. Risks, Limitations, and Validity Threats")
    table(doc, ["Threat to validity", "Current consequence", "Required mitigation"], [
        ["Shared physical model path", "Candidate, commit, acknowledgment, and readback can agree while the common model is wrong", "Use an independently operated model or measurement path and external benchmark review"],
        ["Steady-state abstraction", "Transient, protection, frequency, controller, and hardware behavior are unobserved", "Add fit-for-purpose dynamic, protection, HIL, and device-specific evaluations"],
        ["Guardband selection", "B3 permits three reference-unsafe boundary templates", "Calibrate thresholds and uncertainty against physical-model evidence"],
        ["Synthetic prevalence", "Rates do not estimate incident likelihood", "Report conditional scenario performance only"],
        ["Single host and loopback", "Process behavior is observed, but segmented trust boundaries and host-compromise resilience are untested", "Deploy and negatively test independent services, hosts, identities, and networks"],
        ["Privileged coordinator and shared M4a plant", "Process/key separation can coexist with coordinator compromise or common-mode model error", "Separate administrative domains and use an independently operated measurement or model path"],
        ["Unsigned evidence package", "Internal consistency does not establish origin, custody, or historical integrity", "Use protected signing identity, external anchoring, and documented custody"],
        ["Local reproduction", "Outcome stability is shown only under matching recorded source, lock, host, and component conditions", "Obtain independently operated replication on separately controlled infrastructure"],
        ["Designed condition distribution", "Wilson intervals do not capture model-form or field uncertainty", "Add externally justified operating distributions and sensitivity analysis"],
    ], [1.65, 2.25, 2.6])

    heading(doc, "20. Work Plan and Completion Criteria")
    add_paragraph(doc, "The project advances through controlled governance, executable-kernel, formal-conformance, single-host experiment, physical and PLC integration, multi-node trust-boundary, operate-through-compromise, fleet-scale/economics, and independent-validation gates. The bounded M3 localhost pandapower/PyModbus evaluation, read-only public evidence explorer, and M4a capability-separated deterministic-local implementation are complete local submilestones. WP4 remains in progress because M4a retained evidence, independent sensing or model validation, HELICS coordination, OpenPLC or hardware integration, segmented deployment, durable device identity and replay state, concurrent-controller behavior, post-commit recovery, and external validation have not been completed. A demonstration, green test suite, internally valid evidence package, or perfect fixed-condition result does not satisfy the final completion definition.")

    heading(doc, "21. Conclusions")
    add_paragraph(doc, "The reconstructed foundation now supports two defensible narrow findings and one additional implementation milestone. First, under the reviewed M2 synthetic authorization cases, the complete gateway prevented authorization-invalid execution and outperformed partial control paths, while the separate reference model exposed a material guardband-sensitivity gap. Second, under the five fixed M3 conditions, the signed permit and loopback virtual-device path denied or rejected each registered negative case, completed every nominal command with signed acknowledgment and readback, rejected every replay without a second modeled effect, and preserved the registered narrow proposal/decision/terminal-hash indicator. The local rerun reproduced the deterministic M3 outcome hash under matching recorded conditions. M4a now demonstrates the specified one-dispatch, zero-retry, effect-certainty transaction semantics across distinct same-host plant, signed-observer, and research virtual-PLC processes under local tests and a smoke execution.")
    add_paragraph(doc, "Those results do not establish power-system accuracy, independent device telemetry, retained or independently reproduced M4a evidence, generic Modbus interoperability, segmented trust-boundary enforcement, field failure rates, hard real-time performance, or operational effectiveness. M3 uses one deterministic steady-state model and a Python/PyModbus virtual device whose candidate, commit, and readback paths share a process and model. M4a separates processes and signing identities but still shares the host, coordinator authority, and authoritative plant. The next defensible advance is a retained offline-verifiable M4a package and independently operated consequence or measurement path, followed by separately gated HELICS, OpenPLC or hardware, durable identity and replay state, segmented deployment, recovery, operate-through-compromise, scale, and external validation.")

    heading(doc, "22. References")
    references = [
        "National Institute of Standards and Technology. Guide to Operational Technology (OT) Security. NIST SP 800-82 Rev. 3, September 2023. https://doi.org/10.6028/NIST.SP.800-82r3. Retrieved 2026-08-24.",
        "National Institute of Standards and Technology. Zero Trust Architecture. NIST SP 800-207, August 2020. https://doi.org/10.6028/NIST.SP.800-207. Retrieved 2026-08-24.",
        "National Institute of Standards and Technology. A Zero Trust Architecture Model for Access Control in Cloud-Native Applications in Multi-Cloud Environments. NIST SP 800-207A, September 2023. https://doi.org/10.6028/NIST.SP.800-207A. Retrieved 2026-08-24.",
        "SPIFFE Project. The SPIFFE Standard and Workload API, specification v1.15.2. https://spiffe.io/docs/latest/spiffe-specs/. Retrieved 2026-08-24.",
        "Open Policy Agent. Policy Language and deployment documentation. https://www.openpolicyagent.org/docs/policy-language. Retrieved 2026-08-24.",
        "CIGRE Task Force C6.04. Benchmark Systems for Network Integration of Renewable and Distributed Energy Resources. Technical Brochure 575, 2014. https://www.e-cigre.org/publications/detail/575-benchmark-systems-for-network-integration-of-renewable-and-distributed-energy-resources.html. Retrieved 2026-08-24.",
        "Thurner, L., A. Scheidler, F. Schaefer, J. Menke, J. Dollichon, F. Meier, S. Meinecke, and M. Braun. pandapower--An Open-Source Python Tool for Convenient Modeling, Analysis, and Optimization of Electric Power Systems. IEEE Transactions on Power Systems 33, no. 6 (2018): 6510-6521. https://doi.org/10.1109/TPWRS.2018.2829021.",
        "pandapower. CIGRE Networks, pandapower 3.5.4 documentation. https://pandapower.readthedocs.io/en/develop/networks/cigre.html. Retrieved 2026-08-24.",
        "PyModbus Project. PyModbus v3.15.0 release notes and source repository. https://github.com/pymodbus-dev/pymodbus/releases/tag/v3.15.0. Retrieved 2026-08-24.",
    ]
    add_bullets(doc, references)

    heading(doc, "23. Requirements Traceability Appendix")
    table(doc, ["Requirement", "Implementation", "Test/evidence", "Residual gap"], [
        ["REQ-AUTH-001", "gateway.py; identity.py; delegation.py", "test_unknown_identity; test_valid_full_chain", "No SPIRE service"],
        ["REQ-DELEG-002", "delegation.py", "Full-chain negative tests covering scope, time, depth, signature, expiry, revocation, and linkage", "Mutation testing and distributed revocation remain absent"],
        ["REQ-SAFE-003", "safety.py; oracle.py; pandapower_plant.py", "Independent candidate-state tests; 30-seed M2 comparison; M3 fixed-condition plant results", "No independent physical model or measurement path"],
        ["REQ-REPLAY-004", "replay.py", "Concurrent reservation and retention tests", "Distributed replay store absent"],
        ["REQ-EVID-005", "evidence.py; m3_experiment.py", "Chain integrity, tampering, M3 event correlation, and package verification", "Unsigned manifest; no external anchor"],
        ["REQ-TOCTOU-006", "lab.py; physical_control.py; modbus_device.py", "State-version, candidate re-attestation, compare-and-swap, and restart tests", "No distributed transaction or physical PLC"],
        ["REQ-EXEC-007", "physical_control.py; physical_models.py; modbus_wire.py; modbus_device.py", "Permit, audience-tampering, replay, signed acknowledgment, readback, and offline-verifier tests", "Ephemeral keys and in-memory replay state"],
        ["REQ-CAP-008", "capability_control.py; capability_factory.py; capability_observer.py; capability_plant.py; capability_plc.py", "Capability-negative probes; signature/binding faults; one-dispatch and six-terminal-state tests; orderly child-replacement replay", "Same host and coordinator; no retained M4a package, durable replay, independent sensing, or physical PLC"],
    ], [1.05, 1.8, 2.25, 1.4])

    heading(doc, "24. Data and Message Schemas Appendix")
    add_paragraph(doc, "The ActionProposal schema binds actor, mission, resource, operation, parameters, observed state and time, submission time, nonce, confidence, risk score, delegation chain, and optional approval identifier. Operation-specific parameter sets are closed, and non-finite numeric inputs, out-of-range percentages, extra fields, and timezone-naive timestamps are rejected at the validation boundary. The committed JSON Schema is generated from the authoritative Pydantic model and checked for drift in CI. Decision records bind outcome, reasons, policy and safety versions, state version, decision time, and evidence-record hash.")
    add_paragraph(doc, "M3 exports eight additional closed contracts: physical state, physical command, candidate assessment, execution permit, command acknowledgment, closed-loop result, signed Modbus request, and signed Modbus response. The physical-state contract distinguishes value, topology, observation, model, source, clock, sequence, and state-version bindings. The permit embeds the exact command and binds proposal, decision, candidate, pre-state, expected post-state, policy, safety, evidence, expiry, nonce, signing key, and device audience. The retained package binds the exported schema hashes and the offline verifier also validates current typed runtime contracts.")
    add_paragraph(doc, "M4a exports seven additional closed contracts: capability action request, signed observation envelope, capability execution permit, PLC command acknowledgment, capability closed-loop result, and canonical-JSON IPC request and response frames. Structural schema validation establishes closed fields and internal correlations; live verification separately establishes signatures, trust anchors, freshness, accepted sequence, boot epochs, response counters, and the transaction evidence chain. A structurally valid serialized result is not by itself cryptographic or operational verification.")

    heading(doc, "25. Reproduction Runbook Appendix")
    add_bullets(doc, [
        "Create and activate a clean Python 3.11-or-later virtual environment.",
        "Install with pip install -e \".[dev,docs,simulation]\"; do not rely on PYTHONPATH.",
        "Run python scripts/export_schemas.py --check, python scripts/build_public_demo.py --check, ruff check ., mypy src scripts/build_public_demo.py, and pytest --cov=aegis_ot --cov-branch --cov-report=term-missing --cov-fail-under=90.",
        "Acquire the pinned TLC 1.8.0 JAR, verify SHA-256 eabd140a70f49eb9305a3bd3f3df944eddf87e5a90d329789085f8953a80533a, and run python scripts/run_formal.py --jar /path/to/tla2tools.jar --output-dir results/formal/<run-name>.",
        "Run python -m aegis_ot demo --output-dir results/demo.",
        "Run python -m aegis_ot experiment --trials-per-seed 36 --seed-count 30 --seed 20260824 --output-dir results/m2-independent-oracle.",
        "Run aegis-ot capability-smoke for the bounded M4a local path. Treat its completed status, process identifiers, health counters, one-dispatch/zero-retry values, and in-memory chain check as transient smoke evidence only; the command does not create a retained experiment package.",
        "From a clean committed checkout, run aegis-ot physical-experiment --seed-count 30 --seed 20260824 --output-dir /new/unique/m3-directory; never overwrite a retained package.",
        "Run aegis-ot verify-physical-evidence --output-dir /new/unique/m3-directory from a checkout whose source, schema, project, and lock hashes match the manifest.",
        "Verify clean-start Git state, the raw-data hash, and deterministic outcome hash before reporting results; reproduce into a separate directory and compare outcome hashes.",
        "Build figures and the canonical document from committed scripts, then render and inspect every page.",
    ])

    heading(doc, "26. Version-Control Audit Trail Appendix")
    add_paragraph(doc, "This reconstruction begins a new Git history because the original package and its local uncommitted state were unavailable. No earlier commit hash is represented as an ancestor of this repository. Commit e94e47f records the reconstructed v0.1 foundation, c906ae7 records trust-boundary hardening, 3b5f129 records the bounded formal model and conformance mapping, e8b304d introduces the independent multi-seed evaluator, 050f9b1 corrects clean-start provenance capture, cd20986 adds sampled reviewed scenario envelopes, 168b8bd adds the bounded M3 steady-state plant and signed Modbus process boundary, 0c48c39 retains the verified M3 evidence, 3214b65 adds the evidence-backed read-only public demonstration, and e323bbf adds the bounded M4a capability-separated deterministic-local loop. The controlled M2 evidence and a separate rerun share their deterministic outcome hash. The primary M3 package records clean 168b8bd and a local rerun shares its deterministic outcome hash; those unsigned historical Git fields are internally checked, not externally attested. M4a has local test and smoke evidence but no retained signed-artifact package or offline verifier. No tagged release, external CI run, independent replication, or external validation has been observed.")

    add_header_footer(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    body_text = " ".join(node.text or "" for node in doc.element.body.iter(qn("w:t")))
    words = len(re.findall(r"\b[\w'-]+\b", body_text))
    patch_extended_properties(OUTPUT, int(CURRENT_REVISION["page_count"]), words)
    print(json.dumps({"output": str(OUTPUT), "paragraph_word_count": words, "revision": REVISION_NUMBER}))


if __name__ == "__main__":
    build()
